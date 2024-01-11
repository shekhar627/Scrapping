from urllib.request import urlopen
from bs4 import BeautifulSoup
import docx
from category import getCategories
import json

url = 'https://b.hatena.ne.jp'
allLinksByCategoryDict={}
allLinksByCategoryWithTitleDict={}
categories = getCategories(baseUrl=url)


class LinkWithTitle:
    def __init__(self,title,link,description,subjectList):
        self.title = title
        self.link = link
        self.description = description
        self.subjectList = subjectList

    def toJson(obj):
        return json.dumps(obj, default=lambda obj: obj.__dict__,ensure_ascii=False)

#get all links.
def getLinksWithTitle(content):
    linkWithTitle={}

    soup = BeautifulSoup(content, 'xml')
    allFetchedElements = soup.find_all('item')
    countItem=1
    for element in allFetchedElements:
        _subjectList = {}
        _title = element.find_all('title')[0].contents[0]
        _link = element.find_all('link')[0].contents[0]
        _description = ''
        try:
            _description = element.find_all('description')[0].contents[0]
        except:
            _description=''
        _subjectWithTags = element.find_all('dc:subject')
        countSubject=1
        for subjectElement in _subjectWithTags:
            _subjectList[countSubject] = subjectElement.contents[0]
            countSubject = countSubject+1

        obj = LinkWithTitle(title=_title,link=_link, description=_description,subjectList = _subjectList)
        linkWithTitle[countItem]=obj.toJson()
        countItem = countItem+1
    return linkWithTitle



def fillDictionaryWithData():
    for category in categories:
        content = urlopen(url+category).read()
        linksWithTitle = getLinksWithTitle(content)
        allLinksByCategoryDict[category] = linksWithTitle



def getOnlyLinks():
    doc = docx.Document()
    for key,value in allLinksByCategoryDict.items():
        doc.add_heading(key)
        for obj in value:
            doc.add_paragraph(obj.link)
    doc.save('C:\\Users\\shekhar\\Downloads\\links.docx')

def getLinks_Title():
    doc = docx.Document()
    for key,value in allLinksByCategoryDict.items():
        doc.add_heading(key)
        for obj in value:
            doc.add_paragraph('Title: {}'.format(obj.title))
            doc.add_paragraph('Link: {}'.format(obj.link))
    doc.save('C:\\Users\\shekhar\\Downloads\\links_with_title.docx')



fillDictionaryWithData()

#getOnlyLinks()

#getLinks_Title()

#saving data into a json file.
jsonFile = open('C:\\Users\\shekhar\\Downloads\\object_list.json','w',encoding='utf-16')
json.dump(allLinksByCategoryDict,jsonFile,ensure_ascii=False)




